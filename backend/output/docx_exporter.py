"""
docx_exporter.py — Pure Python DOCX export using python-docx.
No Node.js required.
"""
import io
import re
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _apply_markdown_inline(run_parent, text: str):
    """Add a paragraph with basic **bold** and *italic* markdown rendered."""
    # Split on bold/italic markers
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = run_parent.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = run_parent.add_run(part[1:-1])
            run.italic = True
        else:
            run_parent.add_run(part)


def export_to_docx(topic: str, content: str, doc_type: str = "Research Report") -> bytes:
    """Export content to a properly formatted .docx file. Returns bytes."""
    doc = Document()

    # ── Page margins ──────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Title block ───────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(topic)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # blue

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f"{doc_type}  ·  {date.today().strftime('%B %d, %Y')}")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()  # spacer

    # ── Body content ──────────────────────────────────────
    for line in content.split('\n'):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Heading levels
        if stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=3)
            h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:], level=2)
            h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        elif stripped.startswith('# '):
            h = doc.add_heading(stripped[2:], level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        # Bullet points
        elif stripped.startswith(('- ', '* ', '• ')):
            p = doc.add_paragraph(style='List Bullet')
            _apply_markdown_inline(p, stripped[2:])
        # Numbered list
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            _apply_markdown_inline(p, re.sub(r'^\d+\.\s', '', stripped))
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _apply_markdown_inline(p, stripped)
            for run in p.runs:
                run.font.size = Pt(11)

    # ── Save to bytes ─────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
